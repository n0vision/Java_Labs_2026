# -*- coding: utf-8 -*-
import os
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

BASE = Path(r"c:\Users\Никита\Desktop\java_labs\lab_6")
REPORT_PATH = Path(r"c:\Users\Никита\Desktop\java_labs\Владимиров_Н.А._РИМ-150950_Лаб6_отчет.docx")
GITHUB_URL = "https://github.com/n0vision/Java_Labs_2026"


def add_paragraph(doc, text, size=14, bold=False, align=None, spacing=1.5):
    paragraph = doc.add_paragraph(text)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = spacing
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
    return paragraph


def add_code_block(doc, code):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(10)


def read_java(rel_path):
    return (BASE / rel_path).read_text(encoding="utf-8")


def run_java(class_name, stdin_text=None, demo=False):
    command = [
        "java",
        "-Dfile.encoding=UTF-8",
        "-cp",
        str(BASE / "out"),
    ]
    if demo:
        command.insert(1, "-Dlab6.demo=true")
    command.append(class_name)

    result = subprocess.run(
        command,
        input=stdin_text.encode("utf-8") if stdin_text else None,
        capture_output=True,
        cwd=BASE,
    )
    stdout = result.stdout.decode("utf-8", errors="replace")
    if result.returncode != 0:
        stderr = result.stderr.decode("utf-8", errors="replace")
        raise RuntimeError(f"Java failed for {class_name}: {stderr}")
    return stdout.strip()


def capture_program_output():
    lines = ["(демонстрационный режим: сокращённые задержки для отчёта)", ""]

    lines.append("--- Примеры решения задач ---")
    lines.append("")
    lines.append(run_java("examples_lr6.Example1", demo=True))
    lines.append("")
    lines.append(run_java("examples_lr6.Example2", demo=True))
    lines.append("")
    lines.append(run_java("examples_lr6.Example3"))

    lines.append("")
    lines.append("--- Задания для самостоятельной работы ---")
    lines.append("")
    lines.append(run_java("lr6.task4.Task4"))
    lines.append("")
    lines.append(run_java("lr6.task5.Task5"))
    lines.append("")
    lines.append(run_java("lr6.task6.Task6"))

    lines.append("")
    lines.append("--- Задачи Timus ---")
    lines.append("")
    lines.append("Задача 1082 (сумма чисел от 1 до 10):")
    lines.append(run_java("timus.task_1082.Main"))
    lines.append("")
    lines.append("Задача 1146 (вход: 1 2 3):")
    lines.append(run_java("timus.task_1146.Main", "1 2 3\n"))

    return "\n".join(lines)


def build_report():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    add_paragraph(doc, "Министерство науки и высшего образования Российской Федерации", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Федеральное государственное бюджетное образовательное учреждение", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "высшего образования", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "«Сибирский государственный университет телекоммуникаций и информатики»", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_paragraph(doc, "Кафедра программной инженерии", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "ОТЧЁТ", 14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "по лабораторной работе № 6", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "по дисциплине «Программирование на Java»", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "«Введение в многопоточность»", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "Выполнил: студент группы РИМ-150950", 14)
    add_paragraph(doc, "Владимиров Н.А.", 14)
    doc.add_paragraph()
    add_paragraph(doc, "Проверил: преподаватель кафедры ПИ", 14)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "Новосибирск 2026", 14, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    add_paragraph(doc, "1. Цель работы", 14, bold=True)
    add_paragraph(
        doc,
        "Получить представление о механизме многопоточности в языке программирования Java, "
        "научиться создавать и управлять потоками, использовать методы sleep(), join(), "
        "wait(), notifyAll(), а также применять многопоточность для обработки массивов.",
    )

    add_paragraph(doc, "2. Описание задачи", 14, bold=True)
    add_paragraph(doc, "В разделе 1 методички приведены примеры решения задач:")
    examples = [
        "запуск двух потоков, выводящих имя и текущее время в течение 10 секунд;",
        "запуск потока, выводящего числа от 1 до 10 с задержкой 1 секунда;",
        "два потока выводят чётные и нечётные числа от 1 до 10 с синхронизацией.",
    ]
    for index, example in enumerate(examples, 1):
        add_paragraph(doc, f"{index}. {example}")

    add_paragraph(doc, "В разделе 2 необходимо реализовать задания для самостоятельной работы:")
    tasks = [
        "два потока выводят имя и время (аналог примера 1);",
        "поток выводит числа от 1 до 10 (аналог примера 2);",
        "два потока выводят чётные и нечётные числа (аналог примера 3);",
        "создание 10 потоков, каждый выводит свой номер;",
        "поиск максимального элемента в массиве с помощью многопоточности;",
        "суммирование элементов массива с помощью многопоточности;",
        "решение двух задач с сайта acm.timus.ru.",
    ]
    for index, task in enumerate(tasks, 1):
        add_paragraph(doc, f"{index}. {task}")

    add_paragraph(doc, "3. Ход выполнения", 14, bold=True)
    add_paragraph(
        doc,
        "Создан Java-проект lab_6. Примеры из раздела 1 размещены в пакете examples_lr6 "
        "(классы Example1–Example3). Задания 1–6 реализованы в пакетах lr6.task1–lr6.task6. "
        "Общие функции параллельной обработки массивов вынесены в класс ArrayParallelUtils. "
        "Запуск всех примеров и заданий выполняется через класс Main.",
    )
    add_paragraph(doc, f"Репозиторий с исходным кодом: {GITHUB_URL}")

    add_paragraph(doc, "3.1. Примеры решения задач (раздел 1)", 14, bold=True)
    for index in range(1, 4):
        add_paragraph(doc, f"Пример {index} (класс examples_lr6.Example{index}):", 14, bold=True)
        add_code_block(doc, read_java(f"src/examples_lr6/Example{index}.java"))

    add_paragraph(doc, "3.2. Задания для самостоятельной работы", 14, bold=True)
    for index in range(1, 7):
        add_paragraph(doc, f"Задание {index} (класс lr6.task{index}.Task{index}):", 14, bold=True)
        add_code_block(doc, read_java(f"src/lr6/task{index}/Task{index}.java"))

    add_paragraph(doc, "3.3. Класс ArrayParallelUtils", 14, bold=True)
    add_code_block(doc, read_java("src/lr6/ArrayParallelUtils.java"))

    add_paragraph(doc, "3.4. Класс Main", 14, bold=True)
    add_code_block(doc, read_java("src/lr6/Main.java"))

    add_paragraph(doc, "3.5. Задачи Timus", 14, bold=True)
    add_paragraph(doc, "Задача 1082: вычисление суммы чисел от 1 до 10.")
    add_code_block(doc, read_java("src/timus/task_1082/Main.java"))
    add_paragraph(doc, "Задача 1146 «Суммирование»: сумма целых чисел из входного потока.")
    add_code_block(doc, read_java("src/timus/task_1146/Main.java"))

    add_paragraph(doc, "4. Вывод", 14, bold=True)
    add_paragraph(
        doc,
        "В ходе выполнения лабораторной работы изучены основы многопоточного программирования "
        "в Java. Реализованы примеры создания потоков, синхронизации с помощью ReentrantLock "
        "и Condition, а также параллельная обработка массивов с количеством потоков, равным "
        "числу ядер процессора. Программа успешно компилируется и выполняется. "
        "Дополнительно решены задачи 1082 и 1146 с сайта acm.timus.ru.",
    )

    add_paragraph(doc, "5. Результаты выполнения программы", 14, bold=True)
    add_code_block(doc, capture_program_output())

    doc.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    compile_result = subprocess.run(
        ["cmd", "/c", "compile.bat"],
        cwd=BASE,
        capture_output=True,
    )
    if compile_result.returncode != 0:
        print(compile_result.stdout.decode("cp866", errors="replace"))
        print(compile_result.stderr.decode("cp866", errors="replace"), file=sys.stderr)
        sys.exit(compile_result.returncode)

    os.environ["JAVA_TOOL_OPTIONS"] = "-Dfile.encoding=UTF-8"
    report_path = build_report()
    print(f"Report saved to: {report_path}")
