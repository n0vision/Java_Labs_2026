# -*- coding: utf-8 -*-
import os
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

BASE = Path(r"c:\Users\Никита\Desktop\java_labs\lab_8")
REPORT_PATH = Path(r"c:\Users\Никита\Desktop\java_labs\Владимиров_Н.А._РИМ-150950_Лаб8_отчет.docx")
GITHUB_URL = "https://github.com/n0vision/Java_Labs_2026"
CLASSPATH = f"{BASE / 'out'};{BASE / 'lib' / '*'}"


def add_paragraph(doc, text, size=14, bold=False, align=None, spacing=1.5):
    paragraph = doc.add_paragraph(text)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = spacing
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
    return paragraph


def add_code_block(doc, code):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.0
    lines = code.splitlines() or [""]
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        if index < len(lines) - 1:
            run.add_break()


def decode_process_output(data: bytes) -> str:
    if not data:
        return ""
    if sys.platform == "win32":
        return data.decode("cp1251")
    return data.decode("utf-8")


def normalize_output(text: str) -> str:
    return text.replace(str(BASE), ".").replace(str(BASE).replace("\\", "/"), ".")


def read_java(rel_path):
    return (BASE / rel_path).read_text(encoding="utf-8")


def run_java(class_name, stdin_text=None, demo=False, extra_props=None):
    command = [
        "java",
        "-Dfile.encoding=UTF-8",
        "-cp",
        CLASSPATH,
    ]
    if demo:
        command.insert(1, "-Dlab8.demo=true")
    if extra_props:
        for key, value in extra_props.items():
            command.insert(1, f"-D{key}={value}")
    command.append(class_name)

    result = subprocess.run(
        command,
        input=stdin_text.encode("utf-8") if stdin_text else None,
        capture_output=True,
        cwd=BASE,
    )
    stdout = decode_process_output(result.stdout)
    if result.returncode != 0:
        stderr = decode_process_output(result.stderr)
        return f"[Ошибка выполнения]\n{stderr.strip()}\n{stdout.strip()}".strip()
    return normalize_output(stdout.strip())


def capture_program_output():
    sections = []

    sections.append(
        "Примечание: для заданий 2 и 3 использован демонстрационный режим (-Dlab8.demo=true), "
        "ввод с консоли подставляется автоматически."
    )
    sections.append("")

    sections.append("=== Примеры из раздела 1 (классы Example1–Example8) ===")
    for index in range(1, 9):
        sections.append("")
        sections.append(f">>> java examples_lr8.Example{index}")
        output = run_java(f"examples_lr8.Example{index}")
        if index in (5, 6) and "[Ошибка выполнения]" in output:
            sections.append("(требуется подключение к сети; при офлайн-запуске возможна ошибка)")
        sections.append(output)

    sections.append("")
    sections.append("=== Задание 1. Воспроизведение примеров (lr8.task1.Task1) ===")
    sections.append("")
    sections.append(">>> java lr8.task1.Task1")
    sections.append(run_java("lr8.task1.Task1"))

    sections.append("")
    sections.append("=== Задание 2. XML-парсер (lr8.task2.Task2) ===")
    sections.append(">>> java -Dlab8.demo=true lr8.task2.Task2")
    sections.append(run_java("lr8.task2.Task2", demo=True))

    sections.append("")
    sections.append("=== Задание 3. JSON-парсер (lr8.task3.Task3) ===")
    sections.append(">>> java -Dlab8.demo=true lr8.task3.Task3")
    sections.append(run_java("lr8.task3.Task3", demo=True))

    sections.append("")
    sections.append("=== Задание 4. HTML-парсер (lr8.task4.Task4) ===")
    sections.append(">>> java lr8.task4.Task4")
    sections.append(run_java("lr8.task4.Task4"))

    sections.append("")
    sections.append("=== Задание 5. Excel с обработкой ошибок (lr8.task5.Task5) ===")
    sections.append(">>> java lr8.task5.Task5")
    sections.append(run_java("lr8.task5.Task5"))

    sections.append("")
    sections.append("=== Задачи Timus ===")
    sections.append("")
    sections.append("Задача 1005. Вход: XIV")
    sections.append(">>> java timus.task_1005.Main")
    sections.append(run_java("timus.task_1005.Main", "XIV\n"))
    sections.append("")
    sections.append("Задача 1880. Вход: 2 3 и 5 7")
    sections.append(">>> java timus.task_1880.Main")
    sections.append(run_java("timus.task_1880.Main", "2 3\n5 7\n"))

    return "\n".join(sections)


def build_report():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    add_paragraph(doc, "Министерство науки и высшего образования Российской Федерации", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Федеральное государственное бюджетное образовательное учреждение", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "высшего образования", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "«Сибирский государственный университет телекоммуникаций и информатики»", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_paragraph(doc, "Кафедра программной инженерии", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "ОТЧЁТ", 14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "по лабораторной работе № 8", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "по дисциплине «Программирование на Java»", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "«Введение в работу с XML, JSON, HTML и Excel»", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "Выполнил: студент группы РИМ-150950", 14)
    add_paragraph(doc, "Владимиров Н.А.", 14)
    doc.add_paragraph()
    add_paragraph(doc, "Проверил: преподаватель кафедры ПИ", 14)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "Новосибирск 2026", 14, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    add_paragraph(doc, "1. Цель работы", 14, bold=True)
    add_paragraph(
        doc,
        "Получить навыки работы с форматами обмена данными XML, JSON, HTML и Excel в Java. "
        "Освоить создание и разбор структурированных данных с использованием библиотек JAXP (DOM), "
        "json-simple, jsoup и Apache POI.",
    )

    add_paragraph(doc, "2. Описание задачи", 14, bold=True)
    add_paragraph(doc, "В разделе 1 методички приведены примеры:")
    examples = [
        "создание и чтение XML-файла с помощью DOM;",
        "создание и чтение JSON-файла с помощью json-simple;",
        "парсинг HTML-страниц с помощью jsoup;",
        "создание и чтение Excel-файлов с помощью Apache POI.",
    ]
    for index, example in enumerate(examples, 1):
        add_paragraph(doc, f"{index}. {example}")

    add_paragraph(doc, "В разделе 2 необходимо выполнить задания для самостоятельной работы:")
    tasks = [
        "воспроизвести все примеры из раздела 1;",
        "доработать XML-парсер: создание файла, добавление книг, поиск по автору/году, удаление по названию;",
        "доработать JSON-парсер: поиск по автору, добавление и удаление книг;",
        "доработать HTML-парсер: сохранение данных в файл и обработка ошибок подключения;",
        "улучшить обработку ошибок при чтении Excel-файлов;",
        "решить две задачи с сайта acm.timus.ru.",
    ]
    for index, task in enumerate(tasks, 1):
        add_paragraph(doc, f"{index}. {task}")

    add_paragraph(doc, "Вариант задания: список книг (Book).", 14)

    add_paragraph(doc, "3. Ход выполнения", 14, bold=True)
    add_paragraph(
        doc,
        "Создан Java-проект lab_8. Примеры из раздела 1 размещены в пакете examples_lr8 "
        "(классы Example1–Example8). Задания 1–5 реализованы в пакетах lr8.task1–lr8.task5. "
        "Вспомогательные классы XmlBookService и JsonBookService содержат операции с данными. "
        "Внешние библиотеки подключены через каталог lib. Запуск выполняется через класс Main.",
    )
    add_paragraph(doc, f"Репозиторий с исходным кодом: {GITHUB_URL}")

    add_paragraph(doc, "3.1. Примеры из раздела 1", 14, bold=True)
    for index in range(1, 9):
        add_paragraph(doc, f"Пример {index} (класс examples_lr8.Example{index}):", 14, bold=True)
        add_code_block(doc, read_java(f"src/examples_lr8/Example{index}.java"))

    add_paragraph(doc, "3.2. Задания для самостоятельной работы", 14, bold=True)
    for index in range(1, 6):
        add_paragraph(doc, f"Задание {index} (класс lr8.task{index}.Task{index}):", 14, bold=True)
        add_code_block(doc, read_java(f"src/lr8/task{index}/Task{index}.java"))

    add_paragraph(doc, "3.3. Вспомогательные классы", 14, bold=True)
    add_paragraph(doc, "XmlBookService:", 14, bold=True)
    add_code_block(doc, read_java("src/lr8/XmlBookService.java"))
    add_paragraph(doc, "JsonBookService:", 14, bold=True)
    add_code_block(doc, read_java("src/lr8/JsonBookService.java"))
    add_paragraph(doc, "DataDemoConfig:", 14, bold=True)
    add_code_block(doc, read_java("src/lr8/DataDemoConfig.java"))

    add_paragraph(doc, "3.4. Класс Main", 14, bold=True)
    add_code_block(doc, read_java("src/lr8/Main.java"))

    add_paragraph(doc, "3.5. Задачи Timus", 14, bold=True)
    add_paragraph(doc, "Задача 1005: преобразование римского числа в арабское.")
    add_code_block(doc, read_java("src/timus/task_1005/Main.java"))
    add_paragraph(doc, "Задача 1880 «Снова A+B?!»: сумма пар целых чисел из входного потока.")
    add_code_block(doc, read_java("src/timus/task_1880/Main.java"))

    add_paragraph(doc, "4. Вывод", 14, bold=True)
    add_paragraph(
        doc,
        "В ходе выполнения лабораторной работы изучены способы обработки структурированных данных в Java: "
        "создание и парсинг XML через DOM, работа с JSON при помощи json-simple, извлечение данных из HTML "
        "с помощью jsoup и чтение/запись Excel-файлов через Apache POI. Реализованы все примеры и задания "
        "методички, включая операции добавления, поиска и удаления записей, сохранение результатов парсинга "
        "и обработку ошибок. Программа успешно компилируется и выполняется. "
        "Дополнительно решены задачи 1005 и 1880 с сайта acm.timus.ru.",
    )

    add_paragraph(doc, "5. Результаты выполнения программы", 14, bold=True)
    add_code_block(doc, capture_program_output())

    doc.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    compile_result = subprocess.run(
        ["cmd", "/c", "compile.bat"],
        cwd=BASE,
        capture_output=True,
    )
    if compile_result.returncode != 0:
        print(compile_result.stdout.decode("cp866", errors="replace"))
        print(compile_result.stderr.decode("cp866", errors="replace"), file=sys.stderr)
        sys.exit(compile_result.returncode)

    report_path = build_report()
    print(f"Report saved to: {report_path}")
