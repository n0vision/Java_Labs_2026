# -*- coding: utf-8 -*-
import os
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

BASE = Path(r"c:\Users\Никита\Desktop\java_labs\lab_7")
REPORT_PATH = Path(r"c:\Users\Никита\Desktop\java_labs\Владимиров_Н.А._РИМ-150950_Лаб7_отчет.docx")
GITHUB_URL = "https://github.com/n0vision/Java_Labs_2026"


def add_paragraph(doc, text, size=14, bold=False, align=None, spacing=1.5):
    paragraph = doc.add_paragraph(text)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = spacing
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
    return paragraph


def add_code_block(doc, code):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.0
    lines = code.splitlines() or [""]
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        if index < len(lines) - 1:
            run.add_break()


def decode_process_output(data: bytes) -> str:
    if not data:
        return ""
    if sys.platform == "win32":
        return data.decode("cp1251")
    return data.decode("utf-8")


def normalize_output(text: str) -> str:
    return text.replace(str(BASE), ".").replace(str(BASE).replace("\\", "/"), ".")


def read_java(rel_path):
    return (BASE / rel_path).read_text(encoding="utf-8")


def run_java(class_name, stdin_text=None, demo=False):
    command = [
        "java",
        "-Dfile.encoding=UTF-8",
        "-cp",
        str(BASE / "out"),
    ]
    if demo:
        command.insert(1, "-Dlab7.demo=true")
    command.append(class_name)

    result = subprocess.run(
        command,
        input=stdin_text.encode("utf-8") if stdin_text else None,
        capture_output=True,
        cwd=BASE,
    )
    stdout = decode_process_output(result.stdout)
    if result.returncode != 0:
        stderr = decode_process_output(result.stderr)
        raise RuntimeError(f"Java failed for {class_name}: {stderr}")
    return normalize_output(stdout.strip())


def capture_program_output():
    sections = []

    sections.append(
        "Примечание: для заданий 2, 5, 6 и 7 использован демонстрационный режим (-Dlab7.demo=true), "
        "ввод с консоли подставляется автоматически."
    )
    sections.append("")

    sections.append("=== Примеры из раздела 1 (классы Example1–Example9) ===")
    for index in range(1, 10):
        sections.append("")
        sections.append(f">>> java examples_lr7.Example{index}")
        sections.append(run_java(f"examples_lr7.Example{index}"))

    sections.append("")
    sections.append("=== Задание 1. Воспроизведение примеров (lr7.task1.Task1) ===")
    sections.append("")
    sections.append(">>> java lr7.task1.Task1")
    sections.append(run_java("lr7.task1.Task1"))

    sections.append("")
    sections.append("=== Задание 2. Ввод с консоли (lr7.task2.Task2) ===")
    sections.append("Ввод: demo data from console (5 раз, для каждого типа потока)")
    sections.append("")
    sections.append(">>> java -Dlab7.demo=true lr7.task2.Task2")
    sections.append(run_java("lr7.task2.Task2", demo=True))

    sections.append("")
    sections.append("=== Задания 3–8 ===")
    demo_inputs = {
        5: "файл: task5/size_demo.txt",
        6: "файл: task6/search.txt, слово: Java",
        7: "файл: task7/output.txt, текст: Текст для записи в файл.",
    }
    for index in range(3, 9):
        sections.append("")
        sections.append(f"--- Задание {index} (lr7.task{index}.Task{index}) ---")
        if index in demo_inputs:
            sections.append(f"Ввод: {demo_inputs[index]}")
            sections.append(f">>> java -Dlab7.demo=true lr7.task{index}.Task{index}")
            sections.append(run_java(f"lr7.task{index}.Task{index}", demo=True))
        else:
            sections.append(f">>> java lr7.task{index}.Task{index}")
            sections.append(run_java(f"lr7.task{index}.Task{index}"))

    sections.append("")
    sections.append("=== Задачи Timus ===")
    sections.append("")
    sections.append("Задача 1083. Вход: 1 2 3")
    sections.append(">>> java timus.task_1083.Main")
    sections.append(run_java("timus.task_1083.Main", "1 2 3\n"))
    sections.append("")
    sections.append("Задача 1409. Вход: 2 2 4 1")
    sections.append(">>> java timus.task_1409.Main")
    sections.append(run_java("timus.task_1409.Main", "2 2 4 1\n"))

    return "\n".join(sections)


def build_report():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    add_paragraph(doc, "Министерство науки и высшего образования Российской Федерации", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Федеральное государственное бюджетное образовательное учреждение", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "высшего образования", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "«Сибирский государственный университет телекоммуникаций и информатики»", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_paragraph(doc, "Кафедра программной инженерии", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "ОТЧЁТ", 14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "по лабораторной работе № 7", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "по дисциплине «Программирование на Java»", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "«Система ввода/вывода в Java. Работа с файлами»", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "Выполнил: студент группы РИМ-150950", 14)
    add_paragraph(doc, "Владимиров Н.А.", 14)
    doc.add_paragraph()
    add_paragraph(doc, "Проверил: преподаватель кафедры ПИ", 14)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "Новосибирск 2026", 14, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    add_paragraph(doc, "1. Цель работы", 14, bold=True)
    add_paragraph(
        doc,
        "Получить навыки работы с каталогами и файлами операционной системы, а также с классами "
        "ввода/вывода Java. Научиться выполнять ввод/вывод данных файла через байтовые и символьные потоки, "
        "использовать буферизацию, сериализацию объектов и класс PrintWriter.",
    )

    add_paragraph(doc, "2. Описание задачи", 14, bold=True)
    add_paragraph(doc, "В разделе 1 методички приведены примеры:")
    examples = [
        "работа с классом File;",
        "байтовый ввод/вывод (FileInputStream, FileOutputStream);",
        "символьные потоки (FileReader, FileWriter);",
        "буферизация (BufferedReader, BufferedWriter);",
        "преобразование байтовых потоков в символьные (InputStreamReader, OutputStreamWriter);",
        "запись в файл с помощью PrintWriter;",
        "сериализация объекта Person;",
        "подсчёт строк в файле и копирование файла.",
    ]
    for index, example in enumerate(examples, 1):
        add_paragraph(doc, f"{index}. {example}")

    add_paragraph(doc, "В разделе 2 необходимо выполнить задания для самостоятельной работы:")
    tasks = [
        "воспроизвести все примеры из раздела 1;",
        "доработать примеры для ввода данных с консоли;",
        "подсчитать количество строк в текстовом файле;",
        "скопировать содержимое одного файла в другой;",
        "вывести размер файла в байтах;",
        "найти строки, содержащие заданное слово;",
        "записать текст в файл и вывести количество записанных символов;",
        "сериализовать и десериализовать объект класса Student;",
        "решить две задачи с сайта acm.timus.ru.",
    ]
    for index, task in enumerate(tasks, 1):
        add_paragraph(doc, f"{index}. {task}")

    add_paragraph(doc, "3. Ход выполнения", 14, bold=True)
    add_paragraph(
        doc,
        "Создан Java-проект lab_7. Примеры из раздела 1 размещены в пакете examples_lr7 "
        "(классы Example1–Example9). Задания 1–8 реализованы в пакетах lr7.task1–lr7.task8. "
        "Вспомогательный класс FileDemoConfig задаёт каталог data для демонстрационных файлов. "
        "Запуск примеров и заданий выполняется через класс Main.",
    )
    add_paragraph(doc, f"Репозиторий с исходным кодом: {GITHUB_URL}")

    add_paragraph(doc, "3.1. Примеры из раздела 1", 14, bold=True)
    for index in range(1, 10):
        add_paragraph(doc, f"Пример {index} (класс examples_lr7.Example{index}):", 14, bold=True)
        add_code_block(doc, read_java(f"src/examples_lr7/Example{index}.java"))

    add_paragraph(doc, "3.2. Класс Person (сериализация)", 14, bold=True)
    add_code_block(doc, read_java("src/examples_lr7/Person.java"))

    add_paragraph(doc, "3.3. Задания для самостоятельной работы", 14, bold=True)
    for index in range(1, 9):
        add_paragraph(doc, f"Задание {index} (класс lr7.task{index}.Task{index}):", 14, bold=True)
        add_code_block(doc, read_java(f"src/lr7/task{index}/Task{index}.java"))

    add_paragraph(doc, "3.4. Класс FileDemoConfig", 14, bold=True)
    add_code_block(doc, read_java("src/lr7/FileDemoConfig.java"))

    add_paragraph(doc, "3.5. Класс Main", 14, bold=True)
    add_code_block(doc, read_java("src/lr7/Main.java"))

    add_paragraph(doc, "3.6. Задачи Timus", 14, bold=True)
    add_paragraph(doc, "Задача 1083: сумма целых чисел из входного потока.")
    add_code_block(doc, read_java("src/timus/task_1083/Main.java"))
    add_paragraph(doc, "Задача 1409 «Восьмой аргумент»: вычисление целочисленного частного степеней.")
    add_code_block(doc, read_java("src/timus/task_1409/Main.java"))

    add_paragraph(doc, "4. Вывод", 14, bold=True)
    add_paragraph(
        doc,
        "В ходе выполнения лабораторной работы изучены механизмы ввода/вывода в Java: работа с классом File, "
        "байтовые и символьные потоки, буферизация, адаптеры потоков, PrintWriter и сериализация объектов. "
        "Реализованы все примеры и задания методички, включая операции с текстовыми файлами и сохранение "
        "объектов в бинарном виде. Программа успешно компилируется и выполняется. "
        "Дополнительно решены задачи 1083 и 1409 с сайта acm.timus.ru.",
    )

    add_paragraph(doc, "5. Результаты выполнения программы", 14, bold=True)
    add_code_block(doc, capture_program_output())

    doc.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    compile_result = subprocess.run(
        ["cmd", "/c", "compile.bat"],
        cwd=BASE,
        capture_output=True,
    )
    if compile_result.returncode != 0:
        print(compile_result.stdout.decode("cp866", errors="replace"))
        print(compile_result.stderr.decode("cp866", errors="replace"), file=sys.stderr)
        sys.exit(compile_result.returncode)

    report_path = build_report()
    print(f"Report saved to: {report_path}")
