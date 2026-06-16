# -*- coding: utf-8 -*-
import os
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

BASE = Path(r"c:\Users\Никита\Desktop\java_labs\lab_5")
REPORT_PATH = Path(r"c:\Users\Никита\Desktop\java_labs\Владимиров_Н.А._РИМ-150950_Лаб5_отчет.docx")
GITHUB_URL = "https://github.com/n0vision/Java_Labs_2026"


def add_paragraph(doc, text, size=14, bold=False, align=None, spacing=1.5):
    paragraph = doc.add_paragraph(text)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = spacing
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
    return paragraph


def add_code_block(doc, code):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(10)


def read_java(rel_path):
    return (BASE / rel_path).read_text(encoding="utf-8")


def run_java(class_name, stdin_text=None):
    command = [
        "java",
        "-Dfile.encoding=UTF-8",
        "-cp",
        str(BASE / "out"),
        class_name,
    ]
    result = subprocess.run(
        command,
        input=stdin_text.encode("utf-8") if stdin_text else None,
        capture_output=True,
        cwd=BASE,
    )
    stdout = result.stdout.decode("utf-8", errors="replace")
    if result.returncode != 0:
        stderr = result.stderr.decode("utf-8", errors="replace")
        raise RuntimeError(f"Java failed for {class_name}: {stderr}")
    return stdout.strip()


def capture_program_output():
  lines = [run_java("lr5.Main"), ""]
  lines.append("--- Задачи Timus ---")
  lines.append("")
  lines.append("Задача 1877 (вход: 10, 20):")
  lines.append(run_java("timus.task_1877.Main", "10\n20"))
  lines.append("")
  lines.append("Задача 1263 (вход: 5 чисел 1 2 3 4 5):")
  lines.append(run_java("timus.task_1263.Main", "5\n1 2 3 4 5"))
  return "\n".join(lines)


def build_report():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    add_paragraph(doc, "Министерство науки и высшего образования Российской Федерации", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Федеральное государственное бюджетное образовательное учреждение", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "высшего образования", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "«Сибирский государственный университет телекоммуникаций и информатики»", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_paragraph(doc, "Кафедра программной инженерии", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "ОТЧЁТ", 14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "по лабораторной работе № 5", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "по дисциплине «Программирование на Java»", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "«Введение в функциональное программирование и Stream API»", 14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "Выполнил: студент группы РИМ-150950", 14)
    add_paragraph(doc, "Владимиров Н.А.", 14)
    doc.add_paragraph()
    add_paragraph(doc, "Проверил: преподаватель кафедры ПИ", 14)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "Новосибирск 2026", 14, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    add_paragraph(doc, "1. Цель работы", 14, bold=True)
    add_paragraph(
        doc,
        "Получить представление о функциональном программировании и Stream API в Java, "
        "научиться применять методы filter(), map(), anyMatch() и лямбда-выражения "
        "для обработки массивов и списков.",
    )

    add_paragraph(doc, "2. Описание задачи", 14, bold=True)
    add_paragraph(doc, "В разделе 1 методички приведены примеры решения задач:")
    examples = [
        "фильтрация чётных чисел из массива;",
        "поиск общих элементов двух массивов;",
        "отбор строк, начинающихся с заглавной буквы;",
        "получение списка квадратов чисел.",
    ]
    for index, example in enumerate(examples, 1):
        add_paragraph(doc, f"{index}. {example}")

    add_paragraph(doc, "В разделе 2 необходимо реализовать все примеры и дополнительные задания:")
    tasks = [
        "оставить строки, содержащие заданную подстроку;",
        "оставить числа, делящиеся на заданное число без остатка;",
        "оставить строки длиннее заданного значения;",
        "оставить числа, большие заданного значения;",
        "оставить строки, состоящие только из букв;",
        "оставить числа, меньшие заданного значения;",
        "решить две задачи с сайта acm.timus.ru.",
    ]
    for index, task in enumerate(tasks, 5):
        add_paragraph(doc, f"{index}. {task}")

    add_paragraph(doc, "3. Ход выполнения", 14, bold=True)
    add_paragraph(
        doc,
        "Создан Java-проект lab_5. Примеры из раздела 1 размещены в пакете examples_lr5 "
        "(классы Example1–Example4). Для самостоятельной работы общие функции вынесены "
        "в класс StreamFunctions, демонстрация выполнена в классах Task1–Task10. "
        "Запуск всех примеров и заданий выполняется через класс Main.",
    )
    add_paragraph(doc, f"Репозиторий с исходным кодом: {GITHUB_URL}")

    add_paragraph(doc, "3.1. Примеры решения задач (раздел 1)", 14, bold=True)
    for index in range(1, 5):
        add_paragraph(doc, f"Пример {index} (класс examples_lr5.Example{index}):", 14, bold=True)
        add_code_block(doc, read_java(f"src/examples_lr5/Example{index}.java"))

    add_paragraph(doc, "3.2. Класс StreamFunctions", 14, bold=True)
    add_code_block(doc, read_java("src/lr5/StreamFunctions.java"))

    add_paragraph(doc, "3.3. Задания для самостоятельной работы", 14, bold=True)
    for index in range(1, 11):
        add_paragraph(doc, f"Задание {index} (класс lr5.task{index}.Task{index}):", 14, bold=True)
        add_code_block(doc, read_java(f"src/lr5/task{index}/Task{index}.java"))

    add_paragraph(doc, "3.4. Класс Main", 14, bold=True)
    add_code_block(doc, read_java("src/lr5/Main.java"))

    add_paragraph(doc, "3.5. Задачи Timus", 14, bold=True)
    add_paragraph(doc, "Задача 1877 «HTTP Requests»: сравнение количества GET- и POST-запросов.")
    add_code_block(doc, read_java("src/timus/task_1877/Main.java"))
    add_paragraph(doc, "Задача 1263 «Разминка»: подсчёт количества чётных чисел в последовательности.")
    add_code_block(doc, read_java("src/timus/task_1263/Main.java"))

    add_paragraph(doc, "4. Вывод", 14, bold=True)
    add_paragraph(
        doc,
        "В ходе выполнения лабораторной работы изучены основы функционального программирования "
        "в Java и Stream API. Реализованы примеры из раздела 1 и все задания для самостоятельной "
        "работы с использованием методов filter(), map(), anyMatch() и лямбда-выражений. "
        "Программа успешно компилируется и выполняется. Дополнительно решены задачи 1877 и 1263 "
        "с сайта acm.timus.ru.",
    )

    add_paragraph(doc, "5. Результаты выполнения программы", 14, bold=True)
    add_code_block(doc, capture_program_output())

    doc.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    compile_result = subprocess.run(
        ["cmd", "/c", "compile.bat"],
        cwd=BASE,
        capture_output=True,
    )
    if compile_result.returncode != 0:
        print(compile_result.stdout.decode("cp866", errors="replace"))
        print(compile_result.stderr.decode("cp866", errors="replace"), file=sys.stderr)
        sys.exit(compile_result.returncode)

    os.environ["JAVA_TOOL_OPTIONS"] = "-Dfile.encoding=UTF-8"
    report_path = build_report()
    print(f"Report saved to: {report_path}")
